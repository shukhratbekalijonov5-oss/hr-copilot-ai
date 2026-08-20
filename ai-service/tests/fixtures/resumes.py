"""Fictional resume fixtures.

Every person, employer and project below is invented for testing. No real
candidate data may ever be placed in this repository.

The primary fixture (Ji-woo Han) deliberately HAS evidence of NestJS, Redis
Pub/Sub and production Kubernetes, and deliberately has NO AWS and NO Terraform
experience — those absences are what let the suite prove the system reports
"evidence not found" instead of fabricating a match.

`Kafka` appears in the skills list but nothing describes *Kafka Streams
stateful processing*; that partial match is the canonical NEEDS_HUMAN_REVIEW
case.
"""

from __future__ import annotations

# Facts the requirement-mapping tests assert against. Keeping them beside the
# text makes it obvious when an edit to the resume invalidates a test.
PRESENT_SKILLS = (
    "NestJS", "Redis Pub/Sub", "Kubernetes", "TypeScript", "PostgreSQL", "Go",
)
ABSENT_SKILLS = (
    "AWS", "Terraform", "Salesforce", "Swift",
)

from io import BytesIO

# --- Primary fictional candidate -------------------------------------------

JIWOO_HAN_TEXT = """Ji-woo Han
Backend Engineer
Seoul, South Korea | jiwoo.han@example.test

Summary
Backend engineer with eight years building distributed services for logistics
and marketplace products. Focused on reliability, event-driven design and
developer experience.

Experience
Senior Backend Engineer, Hanwool Logistics (2021 - present)
Built and operated the order orchestration platform using NestJS and
TypeScript. Designed the event fan-out layer on Redis Pub/Sub, which propagates
shipment state changes to six downstream consumers with at-least-once delivery.
Led the migration of the platform to a production Kubernetes cluster running on
bare-metal servers in our Seoul datacentre, including rolling deploys,
horizontal pod autoscaling and pod disruption budgets. Reduced p99 order
placement latency from 820ms to 180ms.

Backend Engineer, Miraen Commerce (2018 - 2021)
Maintained a PostgreSQL-backed catalogue service. Introduced partitioning for
the product events table, which had grown past 400 million rows. Wrote the
internal migration tooling still used by the team.

Projects
Sokdo - an open-source load shedding proxy written in Go. Sheds traffic based
on downstream queue depth rather than fixed rate limits.

Skills
TypeScript, NestJS, Node.js, Go, PostgreSQL, Redis, Redis Pub/Sub, Kubernetes,
Docker, gRPC, Kafka, bare-metal operations

Education
BSc Computer Science, Yonsei University (2014 - 2018)

Languages
Korean (native), English (professional working proficiency)
"""

# --- Second fictional candidate, different organization --------------------

MARCUS_OSEI_TEXT = """Marcus Osei
Data Engineer
Accra, Ghana | marcus.osei@example.test

Summary
Data engineer specialising in batch and streaming pipelines for financial
reporting.

Experience
Data Engineer, Akwaaba Financial (2020 - present)
Built nightly reconciliation pipelines in Python and Apache Spark. Modelled the
warehouse in dbt and BigQuery. Owned data quality alerting.

Skills
Python, Spark, dbt, BigQuery, Airflow, SQL

Education
BSc Statistics, University of Ghana (2016 - 2020)

Languages
English (native), Twi (native)
"""


def build_pdf(text: str, *, title: str = "Resume") -> bytes:
    """Renders text into a real multi-page PDF with a genuine text layer."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    pdf.setTitle(title)

    width, height = LETTER
    margin = inch
    line_height = 14
    y = height - margin

    for raw_line in text.split("\n"):
        # Wrap long lines so the PDF has realistic line breaks.
        for line in _wrap(raw_line, 88):
            if y < margin:
                pdf.showPage()
                y = height - margin
            pdf.setFont("Helvetica", 10)
            pdf.drawString(margin, y, line)
            y -= line_height

    pdf.save()
    return buffer.getvalue()


def build_docx(text: str) -> bytes:
    """Builds a real DOCX with paragraphs and a small skills table."""
    from docx import Document

    document = Document()
    for line in text.split("\n"):
        document.add_paragraph(line)

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Kubernetes"
    table.rows[0].cells[1].text = "Production, 3 years"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- Extraction-artefact fixtures ------------------------------------------

# Emulates per-glyph positioned PDFs as pypdf reads them: one space between
# glyphs of a word, two spaces at real word boundaries.
GLYPH_SPACED_LINES = (
    "R a k h m a t i l l o",
    "A n d r e w",
    "F u l l  S t a c k  D e v e l o p e r",
    "S k i l l s",
    "H T M L ,  C S S ,  S A S S ,  J a v a S c r i p t",
    "a n d r e w 0 3 3 1 r @ g m a i l . c o m",
    "+ 8 2 1 0 5 6 3 7 5 4 2 6",
)


def _render_pdf_lines(
    lines, *, font: str = "Helvetica", size: int = 10, x: int = 72
) -> bytes:
    """Renders exact lines (no wrapping) so spacing artefacts survive."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    y = LETTER[1] - 72
    for line in lines:
        pdf.setFont(font, size)
        pdf.drawString(x, y, line)
        y -= 14
    pdf.save()
    return buffer.getvalue()


def build_letter_spaced_pdf() -> bytes:
    """A PDF whose text layer is glyph-spaced, like per-glyph positioned CVs."""
    return _render_pdf_lines(GLYPH_SPACED_LINES)


def build_multi_column_pdf() -> bytes:
    """Two visual columns, interleaved in content-stream order.

    Plain text-layer readers concatenate rows across the gap; layout-aware
    extraction must keep each column's lines separate.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    left = (
        "Work Experience",
        "Senior Backend Engineer",
        "Hanwool Logistics, Seoul",
        "Built the order orchestration platform",
        "using NestJS and TypeScript.",
    )
    right = ("Skills", "Kubernetes", "Redis Pub/Sub", "PostgreSQL", "Docker")

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    pdf.setFont("Helvetica", 10)
    y = LETTER[1] - 72
    for left_line, right_line in zip(left, right):
        pdf.drawString(72, y, left_line)
        pdf.drawString(380, y, right_line)
        y -= 14
    pdf.save()
    return buffer.getvalue()


KOREAN_LINES = (
    "김민준",
    "백엔드 엔지니어",
    "경력",
    "하늘 물류에서 주문 플랫폼을 개발했습니다.",
    "기술: 쿠버네티스, 레디스",
)


def build_korean_pdf() -> bytes:
    """Korean text (CID font) with one English line mixed in."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    y = LETTER[1] - 72
    for line in KOREAN_LINES:
        pdf.setFont("HYSMyeongJo-Medium", 10)
        pdf.drawString(72, y, line)
        y -= 14
    pdf.setFont("Helvetica", 10)
    pdf.drawString(72, y, "Backend Engineer, Seoul, South Korea")
    pdf.save()
    return buffer.getvalue()


# reportlab's built-in fonts cannot encode Cyrillic; a Unicode TTF from the
# host is used when one exists, and the PDF-level Cyrillic test skips cleanly
# otherwise (string-level Cyrillic coverage does not depend on any font).
_CYRILLIC_FONT_CANDIDATES = (
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/Library/Fonts/Arial.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
)

CYRILLIC_LINES = (
    "Дмитрий Волков",
    "Опытный инженер по данным и аналитике",
    "Навыки: Python, Spark, SQL",
)


def find_cyrillic_font() -> str | None:
    import os

    for path in _CYRILLIC_FONT_CANDIDATES:
        if os.path.exists(path):
            return path
    return None


def build_cyrillic_pdf(font_path: str) -> bytes:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    pdfmetrics.registerFont(TTFont("FixtureCyrillic", font_path))
    return _render_pdf_lines(CYRILLIC_LINES, font="FixtureCyrillic")


def build_empty_pdf() -> bytes:
    """A structurally valid PDF with no text layer (like a scan)."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def build_corrupt_pdf() -> bytes:
    """Correct magic bytes, garbage body."""
    return b"%PDF-1.4\n" + b"\x00\xff" * 200


def _wrap(line: str, width: int) -> list[str]:
    if not line.strip():
        return [""]
    words, out, current = line.split(), [], ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= width:
            current = candidate
        else:
            out.append(current)
            current = word
    if current:
        out.append(current)
    return out
