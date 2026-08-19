"""DOCX text extraction.

python-docx exposes paragraphs and tables but no page boundaries — pagination
is decided by the renderer, not stored in the file. Rather than invent page
numbers, every DOCX block is reported as page 1 and ``page_count`` is 1. A
citation therefore reads "resume.docx" with no page, which is accurate.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.common.errors import CorruptDocumentError, EmptyDocumentError
from app.parsers.base import ParsedDocument, ParsedPage


def parse_docx(data: bytes) -> ParsedDocument:
    try:
        document = Document(BytesIO(data))
    except PackageNotFoundError as exc:
        raise CorruptDocumentError("DOCX could not be read: not a valid package") from exc
    except Exception as exc:
        raise CorruptDocumentError(f"DOCX could not be read: {exc}") from exc

    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    # Skills and language proficiency are often laid out in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                blocks.append(line)

    parsed = ParsedDocument(
        page_count=1,
        pages=[ParsedPage(page_number=1, text="\n".join(blocks))],
    )
    if parsed.is_empty:
        raise EmptyDocumentError("DOCX contains no extractable text")
    return parsed
